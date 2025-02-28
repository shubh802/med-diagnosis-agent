import streamlit as st
from crewai import Agent, Task, Crew
from crewai_tools import ScrapeWebsiteTool, SerperDevTool
from langchain_openai import ChatOpenAI
from dotenv import load_dotenv
from docx import Document
from io import BytesIO
import os
import base64

# Load environment variables
load_dotenv()

# Set API Keys
os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")
os.environ["SERPER_API_KEY"] = os.getenv("SERPER_API_KEY")

# Configure Streamlit page
st.set_page_config(
    page_title="AI Healthcare Assistant",
    page_icon="🩺",
    layout="centered",
    initial_sidebar_state="expanded"
)

# Header Section
st.title("🩺 AI-Powered Healthcare Assistant")
st.markdown(
    """
    Welcome to the **AI Healthcare Assistant**, where cutting-edge technology helps doctors and healthcare providers with preliminary diagnoses and personalized treatment plans. 🚑
    """
)
st.divider()

# Input Section
st.markdown("### 📋 **Patient Information**")
with st.container():
    col1, col2 = st.columns(2)

    with col1:
        gender = st.selectbox('Select Gender:', ('Male', 'Female', 'Other'), help="Choose the patient's gender.")
        age = st.number_input('Enter Age:', min_value=0, max_value=120, value=25, help="Enter the patient's age.")

    with col2:
        symptoms = st.text_area(
            'Enter Symptoms:', 
            placeholder="e.g., fever, cough, headache", 
            help="Describe the symptoms in detail."
        )
        medical_history = st.text_area(
            'Enter Medical History:', 
            placeholder="e.g., diabetes, hypertension", 
            help="Provide details of any relevant medical history."
        )

st.divider()

# Initialize Tools
search_tool = SerperDevTool()
scrape_tool = ScrapeWebsiteTool()

## LLM Model
llm = ChatOpenAI(
    model="gpt-4o-mini",
    temperature=0.1,
    max_tokens=8000
)

# Define Agents
diagnostician = Agent(
    role="Medical Diagnostician",
    goal="Analyze patient symptoms and medical history to provide a preliminary diagnosis.",
    backstory="This agent specializes in diagnosing medical conditions based on patient-reported symptoms and medical history. It uses advanced algorithms and medical knowledge to identify potential health issues.",
    verbose=True,
    allow_delegation=False,
    tools=[search_tool, scrape_tool],
    llm=llm
)

treatment_advisor = Agent(
    role="Treatment Advisor",
    goal="Recommend appropriate treatment plans based on the diagnosis provided by the Medical Diagnostician.",
    backstory="This agent specializes in creating treatment plans tailored to individual patient needs. It considers the diagnosis, patient history, and current best practices in medicine to recommend effective treatments.",
    verbose=True,
    allow_delegation=False,
    tools=[search_tool, scrape_tool],
    llm=llm
)

# Define Tasks
diagnose_task = Task(
    description=(
        "1. Analyze the patient's symptoms ({symptoms}) and medical history ({medical_history}).\n"
        "2. Provide a preliminary diagnosis with possible conditions based on the provided information.\n"
        "3. Limit the diagnosis to the most likely conditions."
    ),
    expected_output="A preliminary diagnosis with a list of possible conditions.",
    agent=diagnostician
)

treatment_task = Task(
    description=(
        "1. Based on the diagnosis, recommend appropriate treatment plans step by step.\n"
        "2. Consider the patient's medical history ({medical_history}) and current symptoms ({symptoms}).\n"
        "3. Provide detailed treatment recommendations, including medications, lifestyle changes, and follow-up care."
    ),
    expected_output="A comprehensive treatment plan tailored to the patient's needs.",
    agent=treatment_advisor
)

# Create Crew
crew = Crew(
    agents=[diagnostician, treatment_advisor],
    tasks=[diagnose_task, treatment_task],
    verbose=True
)

# Generate a Word Document
def generate_docx(result):
    doc = Document()
    doc.add_heading('Healthcare Diagnosis and Treatment Recommendations', level=1)
    doc.add_paragraph(result)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# Generate Download Link
def get_download_link(bio, filename):
    b64 = base64.b64encode(bio.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">📥 Download Diagnosis and Treatment Plan</a>'

# Output Section
if st.button("🧾 Get Diagnosis and Treatment Plan"):
    with st.spinner('🩺 Generating recommendations...'):
        result = crew.kickoff(inputs={"symptoms": symptoms, "medical_history": medical_history})
        
        # Extract relevant data for the document
        if isinstance(result, dict) and 'output' in result:
            diagnosis_and_treatment = result['output']
        else:
            diagnosis_and_treatment = str(result)
        
        # Display results and download link
        st.success("✅ Diagnosis and treatment plan generated successfully!")
        st.markdown("### 📝 **Diagnosis and Treatment Plan**")
        st.write(diagnosis_and_treatment)

        # Generate and display the download link
        docx_file = generate_docx(diagnosis_and_treatment)
        download_link = get_download_link(docx_file, "diagnosis_and_treatment_plan.docx")
        st.markdown(download_link, unsafe_allow_html=True)
